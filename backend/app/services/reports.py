"""Monthly report data builder, DOCX renderer, and orchestration."""

from __future__ import annotations

from dataclasses import replace
from datetime import UTC, datetime
from decimal import ROUND_HALF_UP, Decimal
from io import BytesIO
from uuid import UUID
from zoneinfo import ZoneInfo

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sqlalchemy import or_, select
from sqlalchemy.orm import Session

from app.models.category import Category
from app.models.generated_report import GeneratedReport
from app.models.saving_goal import SavingGoal
from app.models.subscription import Subscription
from app.models.transaction import Transaction
from app.models.user import User
from app.routers._scoping import visible_user_ids
from app.services.envelopes import build_envelope_budget_summary
from app.services.report_charts import render_report_charts
from app.services.report_illustrations import render_report_illustrations
from app.services.report_storage import MinioReportStorage
from app.services.report_types import (
    MonthlyReportData,
    ReportCategoryTotal,
    ReportEnvelopeSummary,
    ReportGoalSummary,
    ReportImage,
    ReportMemberSummary,
    ReportSubscriptionSummary,
)
from app.services.saving_goals import calculate_saving_goal_progress
from app.utils.date_format import format_tr_date
from app.utils.recurrence import monthly_equivalent
from app.utils.tl_format import format_tl

ISTANBUL = ZoneInfo("Europe/Istanbul")
MONEY_QUANT = Decimal("0.01")
PERCENT_QUANT = Decimal("0.1")
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _difference(current: Decimal, previous: Decimal) -> Decimal:
    return _money(current - previous)


def _change_percent(current: Decimal, previous: Decimal) -> Decimal | None:
    if previous == 0:
        return None
    return _percent(current - previous, previous)


def _change_sentence(label: str, current: Decimal, previous: Decimal) -> str:
    diff = _difference(current, previous)
    percent = _change_percent(current, previous)
    if previous == 0:
        if current == 0:
            return f"{label}: Bu dönem ve geçen dönem kayıt yok."
        return f"{label}: Geçen dönem kayıt yokken bu dönem {format_tl(current)} oluştu."
    if diff == 0:
        return f"{label}: Geçen dönemle aynı seviyede, {format_tl(current)}."
    direction = "arttı" if diff > 0 else "azaldı"
    percent_text = f" (%{abs(percent)})" if percent is not None else ""
    return (
        f"{label}: {format_tl(previous)} seviyesinden {format_tl(current)} seviyesine "
        f"{direction}; fark {format_tl(abs(diff))}{percent_text}."
    )


def _status_label(status: str) -> str:
    if status == "over":
        return "Limit aşıldı"
    if status == "watch":
        return "Yakından izlenmeli"
    return "Kontrol altında"


def _goal_status_label(status: str) -> str:
    labels = {
        "on_track": "Planda ilerliyor",
        "at_risk": "Riskli seyrediyor",
        "over_limit": "Limit aşıldı",
        "completed": "Tamamlandı",
    }
    return labels.get(status, "İzleniyor")


def _goal_type_label(goal_type: str) -> str:
    return "Birikim" if goal_type == "accumulation" else "Tasarruf"


def _money(value: Decimal) -> Decimal:
    return value.quantize(MONEY_QUANT, rounding=ROUND_HALF_UP)


def _percent(part: Decimal, total: Decimal) -> Decimal:
    if total == 0:
        return Decimal("0.0")
    return ((part / total) * Decimal("100")).quantize(PERCENT_QUANT, rounding=ROUND_HALF_UP)


def _aware_utc(value: datetime) -> datetime:
    if value.tzinfo is None or value.utcoffset() is None:
        return value.replace(tzinfo=UTC)
    return value.astimezone(UTC)


def _local_month_start(value: datetime) -> datetime:
    local = _aware_utc(value).astimezone(ISTANBUL)
    return local.replace(day=1, hour=0, minute=0, second=0, microsecond=0)


def _add_months(value: datetime, months: int) -> datetime:
    month_index = value.month - 1 + months
    return value.replace(year=value.year + month_index // 12, month=month_index % 12 + 1)


def _periods(now: datetime | None = None) -> tuple[datetime, datetime, datetime, datetime]:
    current_start_local = _local_month_start(now or datetime.now(UTC))
    current_end_local = _add_months(current_start_local, 1)
    previous_start_local = _add_months(current_start_local, -1)
    report_end_local = _aware_utc(now or datetime.now(UTC)).astimezone(ISTANBUL)
    return (
        current_start_local.astimezone(UTC),
        min(report_end_local, current_end_local).astimezone(UTC),
        previous_start_local.astimezone(UTC),
        current_start_local.astimezone(UTC),
    )


def _family_members(db: Session, current_user: User) -> list[User]:
    user_ids = visible_user_ids(current_user)
    return list(
        db.execute(select(User).where(User.id.in_(user_ids)).order_by(User.role.desc(), User.name))
        .scalars()
        .all(),
    )


def _allowed_report_members(db: Session, current_user: User, scope: str) -> list[User]:
    if scope == "family" and current_user.role == "parent":
        return _family_members(db, current_user)
    return [current_user]


def _assessment(data: MonthlyReportData) -> tuple[str, list[str]]:
    if data.total_expense == 0 and data.total_income == 0:
        return (
            "Bu ay için rapora girecek yeterli gelir/gider kaydı yok. İlk kayıtlar eklendikçe koç değerlendirmesi daha netleşir.",
            [
                "Bu ay en az üç gider kategorisini düzenli kaydetmeyi deneyebilirsin.",
                "Ay sonunda kısa bir zarf kontrolü yapmak gelecek raporu güçlendirir.",
            ],
        )
    direction = "arttı" if data.total_expense > data.previous_expense else "azaldı"
    top_category = (
        data.category_totals[0].category_name if data.category_totals else "belirgin bir kategori"
    )
    assessment = (
        f"Bu ay toplam giderin {format_tl(data.total_expense)}. Geçen aya göre gider tarafı "
        f"{direction}; en görünür kalem {top_category}. Net durumun {format_tl(data.total_balance)}."
    )
    suggestions = [
        f"{top_category} için haftalık küçük bir kontrol noktası koymayı düşünebilirsin.",
        "Tekrarlayan ödemeleri ayda bir kez gözden geçirmek bütçeyi sade tutar.",
    ]
    if data.goal_summaries:
        suggestions.append(
            "Aktif hedeflerin için katkı veya kalan limit bilgisini haftalık takip et."
        )
    else:
        suggestions.append(
            "Bir birikim veya kategori tasarruf hedefi açmak bu raporu daha aksiyonlu hale getirir."
        )
    return assessment, suggestions[:3]


def build_monthly_report_data(
    db: Session,
    current_user: User,
    *,
    scope: str = "self",
    include_children: bool = False,
    now: datetime | None = None,
) -> MonthlyReportData:
    """Build a fully scoped monthly report data snapshot."""
    scope_type = "family" if scope == "family" and current_user.role == "parent" else "self"
    if include_children and current_user.role == "parent":
        scope_type = "family"
    members = _allowed_report_members(db, current_user, scope_type)
    user_ids = [member.id for member in members]
    period_start, period_end, previous_start, previous_end = _periods(now)

    categories = list(
        db.execute(
            select(Category).where(or_(Category.user_id.in_(user_ids), Category.user_id.is_(None))),
        )
        .scalars()
        .all(),
    )
    category_names = {category.id: category.name for category in categories}
    transactions = list(
        db.execute(
            select(Transaction).where(
                Transaction.user_id.in_(user_ids),
                Transaction.occurred_at >= previous_start,
                Transaction.occurred_at < period_end,
            ),
        )
        .scalars()
        .all(),
    )

    income = Decimal("0")
    expense = Decimal("0")
    previous_income = Decimal("0")
    previous_expense = Decimal("0")
    category_totals: dict[UUID | None, Decimal] = {}
    category_counts: dict[UUID | None, int] = {}
    member_income = dict.fromkeys(user_ids, Decimal("0"))
    member_expense = dict.fromkeys(user_ids, Decimal("0"))
    member_counts = dict.fromkeys(user_ids, 0)

    for transaction in transactions:
        occurred_at = _aware_utc(transaction.occurred_at)
        amount = Decimal(transaction.amount)
        if period_start <= occurred_at < period_end:
            member_counts[transaction.user_id] += 1
            if transaction.type == "income":
                income += amount
                member_income[transaction.user_id] += amount
            else:
                expense += amount
                member_expense[transaction.user_id] += amount
                category_totals[transaction.category_id] = (
                    category_totals.get(transaction.category_id, Decimal("0")) + amount
                )
                category_counts[transaction.category_id] = (
                    category_counts.get(transaction.category_id, 0) + 1
                )
        elif previous_start <= occurred_at < previous_end:
            if transaction.type == "income":
                previous_income += amount
            else:
                previous_expense += amount

    category_sum = sum(category_totals.values(), Decimal("0"))
    category_rows = [
        ReportCategoryTotal(
            category_name=category_names.get(category_id, "Kategorisiz")
            if category_id is not None
            else "Kategorisiz",
            amount=_money(amount),
            share_percent=_percent(amount, category_sum),
            transaction_count=category_counts.get(category_id, 0),
        )
        for category_id, amount in sorted(
            category_totals.items(), key=lambda item: item[1], reverse=True
        )
    ]
    member_rows = [
        ReportMemberSummary(
            user_id=member.id,
            name=member.name,
            role=member.role,
            age_status=member.age_status,
            income=_money(member_income[member.id]),
            expense=_money(member_expense[member.id]),
            balance=_money(member_income[member.id] - member_expense[member.id]),
            transaction_count=member_counts[member.id],
        )
        for member in members
    ]

    goals = list(
        db.execute(
            select(SavingGoal)
            .where(SavingGoal.user_id.in_(user_ids), SavingGoal.status == "active")
            .order_by(SavingGoal.created_at.desc()),
        )
        .scalars()
        .all(),
    )
    goal_rows: list[ReportGoalSummary] = []
    for goal in goals:
        progress = calculate_saving_goal_progress(db, goal, now=period_end)
        target = progress.goal.target_amount or progress.goal.target_spending_amount
        current: Decimal = (
            progress.goal.current_amount
            if progress.goal.goal_type == "accumulation"
            else progress.actual_spending
        )
        goal_rows.append(
            ReportGoalSummary(
                title=progress.goal.title,
                goal_type=progress.goal.goal_type,
                progress_percent=progress.progress_percent,
                current_amount=_money(current),
                target_amount=_money(target),
                remaining_amount=_money(progress.remaining_amount),
                status_label=progress.status_label,
            ),
        )

    subscriptions = list(
        db.execute(
            select(Subscription)
            .where(Subscription.user_id.in_(user_ids), Subscription.is_active.is_(True))
            .order_by(Subscription.name),
        )
        .scalars()
        .all(),
    )
    subscription_rows: list[ReportSubscriptionSummary] = []
    recurring_income = Decimal("0")
    recurring_expense = Decimal("0")
    for subscription in subscriptions:
        monthly = monthly_equivalent(
            Decimal(subscription.amount),
            subscription.recurrence_interval,
            subscription.recurrence_unit,
            subscription.billing_cycle,
        )
        if subscription.type == "income":
            recurring_income += monthly
        else:
            recurring_expense += monthly
        subscription_rows.append(
            ReportSubscriptionSummary(
                name=subscription.name,
                type=subscription.type,
                monthly_equivalent=_money(monthly),
                next_billing_date=subscription.next_billing_date.isoformat()
                if subscription.next_billing_date is not None
                else None,
            ),
        )

    budget_summary = build_envelope_budget_summary(
        categories=categories,
        current_category_totals=category_totals,
        now=period_end,
    )
    envelope_rows = [
        ReportEnvelopeSummary(
            label=envelope.label,
            budget=_money(envelope.budget),
            spent=_money(envelope.spent),
            remaining=_money(envelope.remaining),
            used_percent=envelope.used_percent,
            status=envelope.status,
            safe_daily_amount=_money(envelope.safe_daily_amount),
            days_left_in_month=envelope.days_left_in_month,
        )
        for envelope in budget_summary.envelopes
        if envelope.budget > 0 or envelope.spent > 0
    ]
    title = "Aile Aylık Koç Raporu" if scope_type == "family" else "Aylık Koç Raporu"
    shell = MonthlyReportData(
        owner_user_id=current_user.id,
        owner_name=current_user.name,
        owner_role=current_user.role,
        owner_finance_level=current_user.finance_level,
        scope_type=scope_type,
        title=title,
        period_start=period_start,
        period_end=period_end,
        previous_period_start=previous_start,
        previous_period_end=previous_end,
        included_user_ids=user_ids,
        included_names=[member.name for member in members],
        total_income=_money(income),
        total_expense=_money(expense),
        total_balance=_money(income - expense),
        previous_income=_money(previous_income),
        previous_expense=_money(previous_expense),
        category_totals=category_rows,
        envelope_summaries=envelope_rows,
        members=member_rows,
        goal_summaries=goal_rows,
        subscriptions=subscription_rows,
        recurring_income_monthly=_money(recurring_income),
        recurring_expense_monthly=_money(recurring_expense),
        budgeted_month=_money(budget_summary.budgeted_month),
        remaining_budget=_money(budget_summary.budgeted_month - expense),
        assessment="",
        suggestions=[],
    )
    assessment, suggestions = _assessment(shell)
    return replace(shell, assessment=assessment, suggestions=suggestions)


def _add_kv_table(document: DocxDocument, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_bullets(document: DocxDocument, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _intro_paragraphs(data: MonthlyReportData) -> list[str]:
    names = ", ".join(data.included_names)
    period = f"{format_tr_date(data.period_start)} - {format_tr_date(data.period_end)}"
    scope_text = "aile kapsamındaki" if data.scope_type == "family" else "kişisel"
    return [
        f"Bu rapor {period} dönemi için {scope_text} gelir, gider, zarf, hedef ve tekrarlayan kayıtlarını özetler. Kapsama giren profil(ler): {names}.",
        "Amaç, sadece toplamları göstermek değil; hangi kalemlerin bütçeyi etkilediğini, hangi alanların izlenmesi gerektiğini ve sonraki ay için hangi küçük aksiyonların işe yarayabileceğini görünür kılmaktır.",
    ]


def _cashflow_notes(data: MonthlyReportData) -> list[str]:
    notes = [
        _change_sentence("Gelir", data.total_income, data.previous_income),
        _change_sentence("Gider", data.total_expense, data.previous_expense),
    ]
    if data.total_income > 0:
        expense_ratio = _percent(data.total_expense, data.total_income)
        notes.append(
            f"Gider/gelir oranı %{expense_ratio}. Bu oran yükseldikçe zarf limitleri ve tekrarlayan giderler daha kritik hale gelir."
        )
    elif data.total_expense > 0:
        notes.append(
            "Bu dönem gider kaydı var ancak gelir kaydı yok. Net durumu doğru okumak için düzenli gelirleri de eklemek faydalı olur."
        )
    if data.total_balance < 0:
        notes.append(
            "Net durum negatif. Bu, dönem içinde giderlerin gelirlerden yüksek göründüğünü gösterir; zorunlu giderler ile esnek giderleri ayrı kontrol etmek iyi olur."
        )
    elif data.total_balance > 0:
        notes.append(
            "Net durum pozitif. Bu fark birikim hedefi, acil durum payı veya gelecek ay zarf tamponu için değerlendirilebilir."
        )
    else:
        notes.append(
            "Net durum başa baş. Küçük bir tampon oluşturmak için en yüksek harcama kategorisinde kontrollü azaltım denenebilir."
        )
    return notes


def _category_notes(data: MonthlyReportData) -> list[str]:
    if not data.category_totals:
        return [
            "Bu dönem kategori bazlı gider bulunmadı. İşlemlere kategori eklemek rapor kalitesini belirgin artırır.",
        ]
    top = data.category_totals[0]
    notes = [
        f"En yüksek gider kategorisi {top.category_name}: {format_tl(top.amount)} ve toplam giderin %{top.share_percent} payı. Bu kategori {top.transaction_count} işlemden oluşuyor.",
    ]
    if len(data.category_totals) >= 3:
        top_three = sum((row.amount for row in data.category_totals[:3]), Decimal("0"))
        notes.append(
            f"İlk üç kategori toplam {format_tl(top_three)} ediyor. Bu alanlar, kısa vadeli bütçe kontrolünde en yüksek kaldıraç noktalarıdır."
        )
    small_rows = [row for row in data.category_totals if row.amount <= Decimal("100.00")]
    if small_rows:
        notes.append(
            "Düşük tutarlı kategoriler ayrı ayrı küçük görünse de ay sonunda tekrarlandığında bütçeyi etkileyebilir; sık tekrarlanan küçük harcamaları not etmek faydalıdır."
        )
    return notes


def _envelope_notes(data: MonthlyReportData) -> list[str]:
    if not data.envelope_summaries:
        return [
            "Bu dönem aktif zarf limiti görünmüyor. Market, fatura, ulaşım veya özel bir zarf açmak sonraki ay güvenli günlük harcama hesabını güçlendirir.",
        ]
    notes: list[str] = []
    watched = [item for item in data.envelope_summaries if item.status in {"watch", "over"}]
    if watched:
        for item in watched[:3]:
            percent = f"%{item.used_percent}" if item.used_percent is not None else "oran yok"
            notes.append(
                f"{item.label}: {format_tl(item.budget)} limitin {format_tl(item.spent)} kısmı kullanıldı ({percent}); durum: {_status_label(item.status)}."
            )
    else:
        notes.append(
            "Zarflar genel olarak kontrol altında görünüyor; limit aşımı veya kritik kullanım sinyali yok."
        )
    positive_remaining = [item for item in data.envelope_summaries if item.remaining > 0]
    if positive_remaining:
        best = max(positive_remaining, key=lambda item: item.remaining)
        notes.append(
            f"En geniş kalan alan {best.label}: {format_tl(best.remaining)}. Ay sonuna kadar yaklaşık günlük güvenli pay {format_tl(best.safe_daily_amount)}."
        )
    if data.remaining_budget < 0:
        notes.append(
            f"Toplam zarf bütçesi {format_tl(abs(data.remaining_budget))} aşılmış görünüyor; gelecek ay limitleri gerçek harcama düzenine göre güncellemek gerekebilir."
        )
    return notes


def _goal_notes(data: MonthlyReportData) -> list[str]:
    if not data.goal_summaries:
        return [
            "Aktif hedef yok. Bir birikim hedefi veya en yüksek gider kategorisi için tasarruf hedefi açmak raporu daha aksiyonlu hale getirir.",
        ]
    notes: list[str] = []
    for goal in data.goal_summaries[:4]:
        notes.append(
            f"{goal.title}: {_goal_type_label(goal.goal_type)} hedefi %{goal.progress_percent} seviyesinde; durum {_goal_status_label(goal.status_label)}, kalan {format_tl(goal.remaining_amount)}."
        )
    if any(goal.status_label in {"at_risk", "over_limit"} for goal in data.goal_summaries):
        notes.append(
            "Riskli hedeflerde haftalık küçük kontrol noktası belirlemek ay sonu sürprizlerini azaltır."
        )
    return notes


def _subscription_notes(data: MonthlyReportData) -> list[str]:
    notes = [
        f"Düzenli gelirlerin aylık etkisi {format_tl(data.recurring_income_monthly)}, düzenli giderlerin aylık etkisi {format_tl(data.recurring_expense_monthly)}.",
    ]
    if data.total_income > 0 and data.recurring_expense_monthly > 0:
        recurring_ratio = _percent(data.recurring_expense_monthly, data.total_income)
        notes.append(
            f"Düzenli gider etkisi bu dönem gelirinin yaklaşık %{recurring_ratio} seviyesinde. Bu oran arttığında iptal/yeniden pazarlık kontrolü öncelik kazanır."
        )
    if data.subscriptions:
        expense_items = [item for item in data.subscriptions if item.type == "expense"]
        income_items = [item for item in data.subscriptions if item.type == "income"]
        notes.append(
            f"Aktif tekrarlayan kayıt sayısı: {len(data.subscriptions)}. Gider kaydı {len(expense_items)}, gelir kaydı {len(income_items)}."
        )
        next_items = [item for item in data.subscriptions if item.next_billing_date]
        if next_items:
            notes.append(
                "Yaklaşan tarihleri olan kayıtları ay başında kontrol etmek nakit akışı planını netleştirir."
            )
    else:
        notes.append(
            "Aktif tekrarlayan kayıt yok. Maaş, kira, fatura veya abonelikleri eklemek gelecek ay projeksiyonlarını güçlendirir."
        )
    return notes


def _next_month_checklist(data: MonthlyReportData) -> list[str]:
    checklist = [
        "Gelir ve giderleri aynı gün içinde kaydet; özellikle satıcı/kaynak alanını boş bırakma.",
        "Ay ortasında en yüksek harcama kategorisini ve zarf durumunu bir kez kontrol et.",
        "Tekrarlayan kayıtların yaklaşan tarihlerini gözden geçir; kullanılmayan giderleri not al.",
    ]
    if data.goal_summaries:
        checklist.append("Aktif hedefler için haftalık katkı veya kalan limit kontrolü yap.")
    else:
        checklist.append(
            "En az bir birikim veya tasarruf hedefi açarak gelecek rapora ölçülebilir hedef ekle."
        )
    if data.scope_type == "family":
        checklist.append(
            "Aile üyeleriyle kısa ve yargısız bir bütçe konuşması yap; odak suçlamak değil, ortak plan çıkarmak olsun."
        )
    return checklist


def _add_image(document: DocxDocument, image: ReportImage) -> None:
    document.add_heading(image.title, level=3)
    stream = BytesIO(image.content)
    document.add_picture(stream, width=Inches(5.8))
    document.add_paragraph(image.description)


def render_docx_report(
    data: MonthlyReportData,
    *,
    charts: list[ReportImage],
    illustrations: list[ReportImage],
) -> bytes:
    """Render the monthly report as a DOCX file."""
    document = Document()
    title = document.add_heading(data.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"Dönem: {format_tr_date(data.period_start)} - {format_tr_date(data.period_end)}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Kapsam: {', '.join(data.included_names)}")
    for paragraph in _intro_paragraphs(data):
        document.add_paragraph(paragraph)

    for image in illustrations[:1]:
        _add_image(document, image)

    document.add_heading("Yönetici Özeti", level=1)
    _add_kv_table(
        document,
        [
            ("Toplam gelir", format_tl(data.total_income)),
            ("Toplam gider", format_tl(data.total_expense)),
            ("Net durum", format_tl(data.total_balance)),
            ("Geçen dönem gelir", format_tl(data.previous_income)),
            ("Geçen dönem gider", format_tl(data.previous_expense)),
            ("Aylık zarf bütçesi", format_tl(data.budgeted_month)),
            ("Kalan zarf bütçesi", format_tl(data.remaining_budget)),
            ("Aktif hedef sayısı", str(len(data.goal_summaries))),
            ("Aktif tekrarlayan kayıt", str(len(data.subscriptions))),
        ],
    )
    document.add_paragraph(data.assessment)

    document.add_heading("Nakit Akışı ve Geçen Ay Karşılaştırması", level=1)
    _add_bullets(document, _cashflow_notes(data))

    document.add_heading("Kategori Dağılımı", level=1)
    if data.category_totals:
        _add_kv_table(
            document,
            [
                (
                    row.category_name,
                    f"{format_tl(row.amount)} · %{row.share_percent} · {row.transaction_count} işlem",
                )
                for row in data.category_totals[:8]
            ],
        )
        _add_bullets(document, _category_notes(data))
    else:
        document.add_paragraph("Bu dönem için kategori bazlı gider bulunamadı.")
        _add_bullets(document, _category_notes(data))

    for image in charts:
        _add_image(document, image)

    document.add_heading("Zarf Bütçesi Durumu", level=1)
    if data.envelope_summaries:
        _add_kv_table(
            document,
            [
                (
                    envelope.label,
                    f"Limit {format_tl(envelope.budget)} · harcanan {format_tl(envelope.spent)} · "
                    f"kalan {format_tl(envelope.remaining)} · {_status_label(envelope.status)}",
                )
                for envelope in data.envelope_summaries[:8]
            ],
        )
    _add_bullets(document, _envelope_notes(data))

    document.add_heading("Tasarruf ve Hedefler", level=1)
    if data.goal_summaries:
        _add_kv_table(
            document,
            [
                (
                    row.title,
                    f"{_goal_type_label(row.goal_type)} · %{row.progress_percent} · "
                    f"mevcut {format_tl(row.current_amount)} / hedef {format_tl(row.target_amount)} · "
                    f"kalan {format_tl(row.remaining_amount)} · {_goal_status_label(row.status_label)}",
                )
                for row in data.goal_summaries[:8]
            ],
        )
    else:
        document.add_paragraph("Aktif birikim veya tasarruf hedefi bulunamadı.")
    _add_bullets(document, _goal_notes(data))

    document.add_heading("Tekrarlayan Gelir/Gider Etkisi", level=1)
    _add_kv_table(
        document,
        [
            ("Düzenli gelir etkisi", format_tl(data.recurring_income_monthly)),
            ("Düzenli gider etkisi", format_tl(data.recurring_expense_monthly)),
        ],
    )
    if data.subscriptions:
        _add_kv_table(
            document,
            [
                (
                    item.name,
                    f"{format_tl(item.monthly_equivalent)} / "
                    f"{'Gelir' if item.type == 'income' else 'Gider'}"
                    + (
                        f" · sonraki tarih {item.next_billing_date}"
                        if item.next_billing_date
                        else ""
                    ),
                )
                for item in data.subscriptions[:8]
            ],
        )
    _add_bullets(document, _subscription_notes(data))

    if data.scope_type == "family":
        document.add_heading("Aile Kırılımı", level=1)
        _add_kv_table(
            document,
            [
                (
                    member.name,
                    f"Gelir {format_tl(member.income)} · gider {format_tl(member.expense)} · "
                    f"net {format_tl(member.balance)} · {member.transaction_count} işlem",
                )
                for member in data.members
            ],
        )
        _add_bullets(
            document,
            [
                "Aile kırılımı karşılaştırma veya suçlama için değil, kimin hangi bütçe alanında desteğe ihtiyaç duyduğunu anlamak içindir.",
                "Çocuk profillerinde harçlık ve küçük harcama alışkanlıkları somut örneklerle konuşulmalıdır.",
            ],
        )

    for image in illustrations[1:]:
        _add_image(document, image)

    document.add_heading("Koç Değerlendirmesi", level=1)
    for suggestion in data.suggestions:
        document.add_paragraph(suggestion, style="List Bullet")

    document.add_heading("Gelecek Ay İçin Kontrol Listesi", level=1)
    _add_bullets(document, _next_month_checklist(data))

    document.add_heading("Notlar ve Sınırlar", level=1)
    document.add_paragraph(
        "Bu rapor yatırım tavsiyesi değildir; bütçe koçluğu ve finansal okuryazarlık amaçlıdır."
    )
    document.add_paragraph(
        "Analiz yalnızca uygulamaya girilen kayıtlar kadar güçlüdür. Eksik gelir, nakit harcama veya kategorisiz işlem varsa sonuçlar gerçek finansal durumun tamamını temsil etmeyebilir."
    )
    document.add_paragraph(
        "Tutarlar Türk lirası üzerinden gösterilir; rapor içindeki grafikler ve tablolar karar desteği sağlar, otomatik finansal karar yerine geçmez."
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_monthly_report_file(
    db: Session,
    current_user: User,
    *,
    scope: str = "self",
    include_children: bool = False,
    include_ai_illustrations: bool = True,
    storage: MinioReportStorage | None = None,
    now: datetime | None = None,
) -> dict[str, object]:
    """Generate and store a private DOCX monthly report."""
    data = build_monthly_report_data(
        db,
        current_user,
        scope=scope,
        include_children=include_children,
        now=now,
    )
    charts = render_report_charts(data)
    illustrations = render_report_illustrations(data, enabled=include_ai_illustrations)
    content = render_docx_report(data, charts=charts, illustrations=illustrations)
    filename = f"cuzdan-kocu-raporu-{data.period_start.astimezone(ISTANBUL).strftime('%Y-%m')}.docx"
    stored = (storage or MinioReportStorage()).save_report(
        user_id=current_user.id,
        filename=filename,
        content=content,
        content_type=DOCX_CONTENT_TYPE,
    )
    report = GeneratedReport(
        user_id=current_user.id,
        scope_type=data.scope_type,
        included_user_ids=[str(user_id) for user_id in data.included_user_ids],
        period_start=data.period_start,
        period_end=data.period_end,
        format="docx",
        title=data.title,
        object_name=stored.object_name,
        content_type=stored.content_type,
        file_size_bytes=stored.file_size_bytes,
        metadata_json={
            "filename": filename,
            "chart_count": len(charts),
            "ai_illustration_count": len(illustrations),
            "scope_type": data.scope_type,
        },
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    return {
        "report_id": str(report.id),
        "title": report.title,
        "format": report.format,
        "filename": filename,
        "download_url": f"/api/reports/{report.id}/download",
        "period_start": report.period_start.isoformat(),
        "period_end": report.period_end.isoformat(),
        "period_start_formatted": format_tr_date(report.period_start),
        "period_end_formatted": format_tr_date(report.period_end),
        "scope_type": report.scope_type,
        "included_names": data.included_names,
        "file_size_bytes": report.file_size_bytes,
        "chart_count": len(charts),
        "ai_illustration_count": len(illustrations),
    }
